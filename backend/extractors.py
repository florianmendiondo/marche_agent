import pdfplumber
from docx import Document
import re
from unidecode import unidecode

def extract_docx(path: str) -> str: # Définition de la fonction
# path: str : chemin vers le fichier .docx à traiter.
# -> str : indique que la fonction renvoie une chaîne de caractères (le texte extrait).

# Ouverture du document Word
    doc = Document(path) # Document(path) charge le fichier Word.
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# doc.paragraphs contient tous les paragraphes du document sous forme d’objets.

# Extraction du texte
#       "\n".join(p.text for p in doc.paragraphs if p.text.strip())
#       On parcourt chaque paragraphe p dans doc.paragraphs.
#       p.text récupère le texte du paragraphe.
#       if p.text.strip() filtre les paragraphes vides (supprime les espaces seuls).
#       "\n".join(...) concatène tous les textes avec un retour à la ligne entre chaque paragraphe.
# 
# Résultat : une seule chaîne de texte contenant tout le contenu non vide du document Word.

def extract_pdf(path: str) -> str:
    # path: str : chemin vers le fichier PDF
    # La fonction retourne une chaîne de texte.
    
    # Initialisation
    text = [] #Liste qui va stocker le texte de chaque page.

# Ouverture du PDF
    with pdfplumber.open(path) as pdf: # pdfplumber.open(path) ouvre le fichier PDF.
        # with garantit que le fichier sera correctement fermé après utilisation.
        for page in pdf.pages: # pdf.pages contient toutes les pages du PDF.
            page_text = page.extract_text() # page.extract_text() récupère le texte de la page.
            if page_text: # Si la page contient du texte (if page_text), on l’ajoute à la liste text.
                text.append(page_text) # Concaténation
    return "\n".join(text) ## On assemble tout le texte des pages en une seule chaîne, séparée par des retours à la ligne.
